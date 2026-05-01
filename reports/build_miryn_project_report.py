from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_PATH = ROOT / "miryn_project_report_source.md"
OUTPUT_PATH = ROOT / "Miryn_AI_Project_Report_PhD_Level.docx"


def set_run_font(run, name: str, size: float, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, "Garamond", 10.5)


def strip_md(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Garamond"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Garamond")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Garamond")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)

    for style_name, size, bold in [
        ("Title", 24, True),
        ("Subtitle", 16, False),
        ("Heading 1", 16, True),
        ("Heading 2", 13.5, True),
        ("Heading 3", 12.5, True),
    ]:
        if style_name not in styles and style_name == "Subtitle":
            subtitle = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
            subtitle.base_style = styles["Normal"]
        style = styles[style_name]
        style.font.name = "Garamond"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Garamond")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Garamond")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_after = Pt(8 if "Heading" in style_name else 12)
        if style_name.startswith("Heading"):
            style.paragraph_format.space_before = Pt(12)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)


def add_title_page(doc: Document, source_text: str):
    lines = source_text.splitlines()
    title = strip_md(lines[0].lstrip("# ").strip())
    subtitle = strip_md(lines[2].lstrip("## ").strip())

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, "Garamond", 24, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, "Garamond", 15.5, italic=True)

    for spacer in [0.6, 0.6, 0.6]:
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    for line in lines[4:7]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(strip_md(line))
        set_run_font(run, "Garamond", 12.5)

    doc.add_page_break()


def collect_headings(source_text: str):
    headings = []
    for line in source_text.splitlines():
        if line.startswith("## "):
            headings.append((1, strip_md(line[3:])))
        elif line.startswith("### "):
            headings.append((2, strip_md(line[4:])))
    return headings


def add_contents_page(doc: Document, headings):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Table of Contents")

    for level, text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2 if level == 1 else 0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, "Garamond", 11.5)

    doc.add_page_break()


def add_code_block(doc: Document, block_lines):
    for line in block_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line.rstrip())
        set_run_font(run, "Consolas", 9.5)


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2 + (0.2 * level))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(strip_md(text))
    set_run_font(run, "Garamond", 12)


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(strip_md(text))
    set_run_font(run, "Garamond", 12)


def add_paragraph_text(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(strip_md(text))
    set_run_font(run, "Garamond", 12)


def build_body(doc: Document, source_text: str):
    lines = source_text.splitlines()
    body_lines = lines[8:]
    i = 0
    in_code = False
    code_lines = []
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        if text:
            add_paragraph_text(doc, text)
        paragraph_buffer = []

    while i < len(body_lines):
        raw = body_lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            flush_paragraph()
            i += 1
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(strip_md(line[3:]))
            set_run_font(run, "Garamond", 16, bold=True)
            i += 1
            continue

        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(strip_md(line[4:]))
            set_run_font(run, "Garamond", 13.5, bold=True)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            flush_paragraph()
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue

        if line.startswith("- "):
            flush_paragraph()
            add_bullet(doc, line[2:])
            i += 1
            continue

        paragraph_buffer.append(line)
        i += 1

    flush_paragraph()


def build_report():
    source_text = SOURCE_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_title_page(doc, source_text)
    add_contents_page(doc, collect_headings(source_text))
    build_body(doc, source_text)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
